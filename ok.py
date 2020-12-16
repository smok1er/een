from os import system
from os import unlink
from PyPDF2 import PdfFileReader, PdfFileWriter
from gtts import gTTS
import pytesseract
try:
    from selenium import webdriver
except:
    system('pip install selenium')
import telebot
from docx import Document
from telebot import types
from googletrans import Translator
import requests
from PIL import Image
import io
from PIL import Image
import pytesseract
from wand.image import Image as wi
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'


r1 = requests.session()
translator = Translator()
idi = 209501902
bot = telebot.TeleBot("1260271179:AAEkXs7bInqOvDWUjjJnZuFuqaHI8SfAu9Y")
open('tran.txt', 'a').write('')
c = []


@bot.message_handler(commands=['start'])
def key(msg):
    ch = msg.chat.id
    if ch in c:
        pass
    else:
        c.append(ch)
        bot.send_video(ch, 'BAACAgIAAxkBAAIL7F_TexepMJtQKwNs2WMzTr7vAAEHSgACwggAArfAmEowcyLRDyG4BB4E')
        bot.send_voice(ch, 'AwACAgIAAxkBAAIXPF_V4bRzq7gYQk3CVDMa2tdlzhg2AAIqCQAC4bKxSjIYddAucA69HgQ')
    if msg.chat.id == idi:
        admin = types.ReplyKeyboardMarkup()
        akoky = types.KeyboardButton('الاذاعة')
        zkoky = types.KeyboardButton('الاعضاء')
        admin.add(akoky, zkoky)
        bot.reply_to(msg, 'اهلا و سهلا بك ايها المطور {}'.format(msg.from_user.first_name), reply_markup=admin)

    else:
        q = types.InlineKeyboardMarkup()
        q1 = types.InlineKeyboardButton('اضغط هنا \nاذا كان هنالك مصطلح طبي تريد ترجمته', callback_data='q1')
        a9 = types.InlineKeyboardButton("المطور 💚", url="https://t.me/Q5QQQQ")
        q.add(q1)
        q.add(a9)
        bot.send_message(msg.chat.id, 'اهلا و سهلا بك في بالبوت الترجمة 💚', reply_markup=q)
        if f'{msg.chat.id}\n' in open('tran.txt', 'r'):
            pass
        else:
            open('tran.txt', 'a').write(f'{msg.chat.id}\n')


@bot.message_handler(content_types=['photo'])
def key(msg):
    try:
        ch = msg.chat.id
        file_info = bot.get_file(
            str(msg).split("file_id': '")[-1][:str(msg).split("file_id': '")[-1].find("', 'file_unique_id")])
        downloaded_file = bot.download_file(file_info.file_path)
        with open('xc.jpg', 'wb') as new_file:
            new_file.write(downloaded_file)
        im = Image.open('xc.jpg')
        text = pytesseract.image_to_string(im)
        url = 'https://www.arabtran.com/gtranslate/'

        head = {
        'accept': '*/*',
        'accept-encoding': 'gzip, deflate, br',
        'accept-language': 'en-US,en;q=0.9',
        'content-length': '31',
        'content-type': 'application/x-www-form-urlencoded; charset=UTF-8',
        'cookie': '_ga=GA1.2.189101859.1607603223; _gid=GA1.2.154684385.1607603223; __gads=ID=ba2b5a0fe52a4c1b-22c3694088a60095:T=1607603223:RT=1607603223:S=ALNI_MbdZ0H2MeiuATaH2_vAl4hrUnPm8Q; _gat=1; sc_is_visitor_unique=rx12068393.1607603329.CA2EB2FDB25A4FB19C08D5DA0B092EBC.1.1.1.1.1.1.1.1.1',
        'origin': 'https://www.arabtran.com',
        'referer': 'https://www.arabtran.com/tarjamat_anjilizi_earabiun/',
        'sec-fetch-dest': 'empty',
        'sec-fetch-mode': 'cors',
        'sec-fetch-site': 'same-origin',
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/87.0.4280.88 Safari/537.36',
        'x-requested-with': 'XMLHttpRequest',
        }
        data = {
            'text': text,
            'gfrom': 'en',
            'gto': 'ar',
            'key': 'ABC'
        }
        x = r1.post(url=url, data=data, ).text
        bot.send_message(ch, x)
        tss = gTTS(text=text.replace(' ', '      '), lang='en')
        tss.save('ahmed.mp3')
        bot.send_audio(ch, open('ahmed.mp3', 'rb'), caption='صوت للنصوص الانكليزية الموجودة في الصورة التي ارسلتها 💚')
        bot.send_message(idi, f'@{msg.from_usre.username}')
    except:
        pass


@bot.message_handler(content_types='text')
def an(msg):
    try:
        if msg.text == 'الاعضاء' and msg.chat.id == idi:
            x = open('tran.txt', 'r').readlines()
            bot.reply_to(msg, '{}'.format(len(x)))
    except:
        pass
    try:
        if msg.text == 'الاذاعة' and msg.chat.id == idi:
            markup = types.ForceReply(selective=False)
            bot.send_message(msg.chat.id, "ارسل اذاعتك", reply_markup=markup)
    except:
        pass
    try:
        if msg.reply_to_message.text == "ارسل اذاعتك":
            s = msg.text
            x = open('tran.txt', 'r').readlines()
            z = 0
            try:
                while z < len(x):
                    if str(x[z]) == '\n':
                        pass
                    else:
                        bot.send_message(x[z], s)
                    z += 1
            except:
                pass
    except:
        pass
    w = msg.text.lower()
    try:
        if msg.reply_to_message.text == "ارسل المصطلح الطبي المراد ترجمته":
            if w[0] == 'ا' or w[0] == 'أ' and w[1] == 'ل':
                url = f'https://context.reverso.net/translation/arabic-english/{w}'
                head = {
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
                    'Accept-Encoding': 'gzip, deflate, br',
                    'Accept-Language': 'en-US,en;q=0.9',
                    'Connection': 'keep-alive',
                    'Cookie': 'didomi_token=eyJ1c2VyX2lkIjoiMTc1YWYxY2QtM2E4NS02NWNhLTlhNWYtNWMzMTMxYmY0OTQyIiwiY3JlYXRlZCI6IjIwMjAtMTEtMDlUMjI6MjU6MDMuODM3WiIsInVwZGF0ZWQiOiIyMDIwLTExLTA5VDIyOjI1OjAzLjgzN1oiLCJ2ZXJzaW9uIjpudWxsfQ==; __qca=P0-2069518181-1604960703577; _ga=GA1.2.1022455205.1604960704; _fbp=fb.1.1604960704785.1696161273; __gads=ID=6a415d82c09e7f2f:T=1604961014:S=ALNI_MYnjNMXo7nFRuuahoJIpqXamOQFvA; CTXTNODEID=bstweb12; experiment_context_N7gT3vKzX=0; JSESSIONID=kamXv4rTSPJG1mzG5D0fBsyY.bst-web12; context.lastpair=ar-en; history_entry=psycho]#[gastric]#[{w}; history_pair=en-ar]#[en-ar]#[ar-en; experiment_context_E3de3pqAZ=1; _gid=GA1.2.1607272187.1606403162; context.dapppromotion-count2=1; context.dapppromotion2=0; _gat_gtag_UA_2834324_41=1',
                    'Host': 'context.reverso.net',
                    'Referer': 'https://context.reverso.net/translation/arabic-english/%7Bw%7D',
                    'Sec-Fetch-Dest': 'document',
                    'Sec-Fetch-Mode': 'navigate',
                    'Sec-Fetch-Site': 'same-origin',
                    'Sec-Fetch-User': '?1',
                    'Upgrade-Insecure-Requests': '1',
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36'
                }
                j = r1.get(url=url, headers=head)
                v = (j.text[j.text.find('<button class="other-content" data-other="0" data-negative="') + len(
                    '<button class="other-content" data-other="0" data-negative="'):j.text.find(
                    'Other translations</button>')])[:-2]
                c = (v.replace("-{", "").replace("}", "\n"))
                bot.send_message(msg.chat.id, f'* {c} *')

            else:
                try:
                    if f'{w[0]}' in ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p',
                                     'q',
                                     'r',
                                     's', 't', 'u', 'v', 'w', 'x', 'y', 'z']:
                        url = f'https://context.reverso.net/translation/arabic-english/{w}'
                        head = {
                            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
                            'Accept-Encoding': 'gzip, deflate, br',
                            'Accept-Language': 'en-US,en;q=0.9',
                            'Connection': 'keep-alive',
                            'Cookie': 'didomi_token=eyJ1c2VyX2lkIjoiMTc1YWYxY2QtM2E4NS02NWNhLTlhNWYtNWMzMTMxYmY0OTQyIiwiY3JlYXRlZCI6IjIwMjAtMTEtMDlUMjI6MjU6MDMuODM3WiIsInVwZGF0ZWQiOiIyMDIwLTExLTA5VDIyOjI1OjAzLjgzN1oiLCJ2ZXJzaW9uIjpudWxsfQ==; __qca=P0-2069518181-1604960703577; _ga=GA1.2.1022455205.1604960704; _fbp=fb.1.1604960704785.1696161273; __gads=ID=6a415d82c09e7f2f:T=1604961014:S=ALNI_MYnjNMXo7nFRuuahoJIpqXamOQFvA; CTXTNODEID=bstweb12; experiment_context_N7gT3vKzX=0; JSESSIONID=kamXv4rTSPJG1mzG5D0fBsyY.bst-web12; context.lastpair=ar-en; history_entry=psycho]#[gastric]#[{w}; history_pair=en-ar]#[en-ar]#[ar-en; experiment_context_E3de3pqAZ=1; _gid=GA1.2.1607272187.1606403162; context.dapppromotion-count2=1; context.dapppromotion2=0; _gat_gtag_UA_2834324_41=1',
                            'Host': 'context.reverso.net',
                            'Referer': 'https://context.reverso.net/translation/arabic-english/%7Bw%7D',
                            'Sec-Fetch-Dest': 'document',
                            'Sec-Fetch-Mode': 'navigate',
                            'Sec-Fetch-Site': 'same-origin',
                            'Sec-Fetch-User': '?1',
                            'Upgrade-Insecure-Requests': '1',
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36'
                        }
                        j = r1.get(url=url, headers=head)
                        v = (j.text[j.text.find('<button class="other-content" data-other="0" data-negative="') + len(
                            '<button class="other-content" data-other="0" data-negative="'):j.text.find(
                            'Other translations</button>')])[:-2]
                        c = (v.replace("-{", "").replace("}", "\n"))
                        bot.send_message(msg.chat.id, f'* {c} *')

                    else:
                        url = f'https://context.reverso.net/translation/arabic-english/{w}'
                        head = {
                            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
                            'Accept-Encoding': 'gzip, deflate, br',
                            'Accept-Language': 'en-US,en;q=0.9',
                            'Connection': 'keep-alive',
                            'Cookie': 'didomi_token=eyJ1c2VyX2lkIjoiMTc1YWYxY2QtM2E4NS02NWNhLTlhNWYtNWMzMTMxYmY0OTQyIiwiY3JlYXRlZCI6IjIwMjAtMTEtMDlUMjI6MjU6MDMuODM3WiIsInVwZGF0ZWQiOiIyMDIwLTExLTA5VDIyOjI1OjAzLjgzN1oiLCJ2ZXJzaW9uIjpudWxsfQ==; __qca=P0-2069518181-1604960703577; _ga=GA1.2.1022455205.1604960704; _fbp=fb.1.1604960704785.1696161273; __gads=ID=6a415d82c09e7f2f:T=1604961014:S=ALNI_MYnjNMXo7nFRuuahoJIpqXamOQFvA; CTXTNODEID=bstweb12; experiment_context_N7gT3vKzX=0; JSESSIONID=kamXv4rTSPJG1mzG5D0fBsyY.bst-web12; context.lastpair=ar-en; history_entry=psycho]#[gastric]#[{w}; history_pair=en-ar]#[en-ar]#[ar-en; experiment_context_E3de3pqAZ=1; _gid=GA1.2.1607272187.1606403162; context.dapppromotion-count2=1; context.dapppromotion2=0; _gat_gtag_UA_2834324_41=1',
                            'Host': 'context.reverso.net',
                            'Referer': 'https://context.reverso.net/translation/arabic-english/%7Bw%7D',
                            'Sec-Fetch-Dest': 'document',
                            'Sec-Fetch-Mode': 'navigate',
                            'Sec-Fetch-Site': 'same-origin',
                            'Sec-Fetch-User': '?1',
                            'Upgrade-Insecure-Requests': '1',
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36'
                        }
                        j = r1.get(url=url, headers=head)
                        v = (j.text[j.text.find('<button class="other-content" data-other="0" data-negative="') + len(
                            '<button class="other-content" data-other="0" data-negative="'):j.text.find(
                            'Other translations</button>')])[:-2]
                        c = (v.replace("-{", "").replace("}", "\n"))
                        bot.send_message(msg.chat.id, f'* {c} *')
                except:
                    pass

    except:
        pass
    if msg.reply_to_message:
        pass
    else:
        try:
            url = 'https://www.arabtran.com/gtranslate/'

            head = {
                'accept': '*/*',
                'accept-encoding': 'gzip, deflate, br',
                'accept-language': 'en-US,en;q=0.9',
                'content-length': '31',
                'content-type': 'application/x-www-form-urlencoded; charset=UTF-8',
                'cookie': '_ga=GA1.2.189101859.1607603223; _gid=GA1.2.154684385.1607603223; __gads=ID=ba2b5a0fe52a4c1b-22c3694088a60095:T=1607603223:RT=1607603223:S=ALNI_MbdZ0H2MeiuATaH2_vAl4hrUnPm8Q; _gat=1; sc_is_visitor_unique=rx12068393.1607603329.CA2EB2FDB25A4FB19C08D5DA0B092EBC.1.1.1.1.1.1.1.1.1',
                'origin': 'https://www.arabtran.com',
                'referer': 'https://www.arabtran.com/tarjamat_anjilizi_earabiun/',
                'sec-fetch-dest': 'empty',
                'sec-fetch-mode': 'cors',
                'sec-fetch-site': 'same-origin',
                'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/87.0.4280.88 Safari/537.36',
                'x-requested-with': 'XMLHttpRequest',
            }


            q = types.InlineKeyboardMarkup()
            q1 = types.InlineKeyboardButton('اضغط هنا في حيال وجود مصطلح طبي 💟', callback_data='q1')
            q.add(q1)
            if f'{w[0]}' in ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r',
                             's', 't', 'u', 'v', 'w', 'x', 'y', 'z'] or f'{w[1]}' in ['a', 'b', 'c', 'd', 'e', 'f', 'g',
                                                                                      'h', 'i', 'j', 'k', 'l', 'm', 'n',
                                                                                      'o', 'p', 'q', 'r', 's', 't', 'u',
                                                                                      'v', 'w', 'x', 'y',
                                                                                      'z'] or f'{w[2]}' in ['a', 'b',
                                                                                                            'c', 'd',
                                                                                                            'e', 'f',
                                                                                                            'g', 'h',
                                                                                                            'i', 'j',
                                                                                                            'k', 'l',
                                                                                                            'm', 'n',
                                                                                                            'o', 'p',
                                                                                                            'q', 'r',
                                                                                                            's', 't',
                                                                                                            'u', 'v',
                                                                                                            'w', 'x',
                                                                                                            'y', 'z']:
                data = {
                    'text': w,
                    'gfrom': 'en',
                    'gto': 'ar',
                    'key': 'ABC'
                }
                x = r1.post(url=url, data=data, ).text
                bot.send_message(msg.chat.id, x, reply_markup=q)
            else:
                data = {
                    'text': w,
                    'gfrom': 'ar',
                    'gto': 'en',
                    'key': 'ABC'
                }
                x = r1.post(url=url, data=data, ).text
                bot.send_message(msg.chat.id, x, reply_markup=q)
        except:
            pass


@bot.callback_query_handler(lambda call: True)
def any(call):
    if call.data == 'q1':
        bot.delete_message(call.message.chat.id, call.message.message_id)
        markup = types.ForceReply(selective=False)
        bot.send_message(call.message.chat.id, "ارسل المصطلح الطبي المراد ترجمته", reply_markup=markup)
h = []


@bot.message_handler(content_types='document')
def an(msg):
    ch = msg.chat.id

    open(f'userstran/{ch}.txt', 'w').write('')
    try:
        file_info = bot.get_file(msg.document.file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        x = msg.document.file_name[-4:]
        if x == '.pdf':
            if ch in h:
                bot.send_message(ch, 'انتضر قليلا من فضلك لكي يتم التحميل 💚')
            else:
                h.append(ch)
                bot.send_message(ch, 'حسنا, سيتم التحميل الرجاء انتظار دقيقة واحدة 💚')
                with open('ahmed.pdf', 'wb') as new_file:
                    new_file.write(downloaded_file)
                pdf = wi(filename="ahmed.pdf", resolution=75)
                pdfImage = pdf.convert('jpeg')

                imageBlobs = []

                for img in pdfImage.sequence:
                    imgPage = wi(image=img)
                    imageBlobs.append(imgPage.make_blob('jpeg'))

                recognized_text = []

                for imgBlob in imageBlobs:
                    im = Image.open(io.BytesIO(imgBlob))
                    text = pytesseract.image_to_string(im, lang='eng')
                    recognized_text.append(text)
                w = ''
                for k in recognized_text:
                    try:
                        w += k
                        url = 'https://www.arabtran.com/gtranslate/'

                        head = {
                            'accept': '*/*',
                            'accept-encoding': 'gzip, deflate, br',
                            'accept-language': 'en-US,en;q=0.9',
                            'content-length': '31',
                            'content-type': 'application/x-www-form-urlencoded; charset=UTF-8',
                            'cookie': '_ga=GA1.2.189101859.1607603223; _gid=GA1.2.154684385.1607603223; __gads=ID=ba2b5a0fe52a4c1b-22c3694088a60095:T=1607603223:RT=1607603223:S=ALNI_MbdZ0H2MeiuATaH2_vAl4hrUnPm8Q; _gat=1; sc_is_visitor_unique=rx12068393.1607603329.CA2EB2FDB25A4FB19C08D5DA0B092EBC.1.1.1.1.1.1.1.1.1',
                            'origin': 'https://www.arabtran.com',
                            'referer': 'https://www.arabtran.com/tarjamat_anjilizi_earabiun/',
                            'sec-fetch-dest': 'empty',
                            'sec-fetch-mode': 'cors',
                            'sec-fetch-site': 'same-origin',
                            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/87.0.4280.88 Safari/537.36',
                            'x-requested-with': 'XMLHttpRequest',
                        }

                        data = {
                            'text': k,
                            'gfrom': 'en',
                            'gto': 'ar',
                            'key': 'ABC'
                        }
                        x = r1.post(url=url, data=data, ).text
                        open(f'userstran/{ch}.txt', 'a', encoding='utf-8').write(f"{x}{k}\n")

                    except:
                        pass
                do = Document()
                file = open(f'userstran/{ch}.txt', 'r', encoding='utf-8').read()

                do.add_heading(file, level=1)
                x = do.add_paragraph()
                do.save('Ahmed.docx')
                bot.send_document(ch, open('Ahmed.docx', 'rb'), caption=msg.document.file_name)
                h.remove(ch)
                bot.send_message(ch, 'انتضر قليلا لتلقي الصوت💚 ')
                tss = gTTS(text=w.replace(' ', '    '), lang='en')
                tss.save('ahmed.mp3')
                bot.send_audio(ch, open('ahmed.mp3', 'rb'), caption='صوت للنصوص الانكليزية الموجودة في ال pdf الذي ارسلته 💚')
        else:
            bot.send_message(ch, 'من فضلك ارسل الملفات بصيغ ال pdf ')
    except:
        pass


bot.polling()

