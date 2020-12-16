from os import system
from os import unlink
from PyPDF2 import PdfFileReader, PdfFileWriter
from gtts import gTTS
import pytesseract
import cv2
try:
    from selenium import webdriver
except:
    system('pip install selenium')
import telebot
from docx import Document
from telebot import types
from googletrans import Translator
import requests
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
r1 = requests.session()
translator = Translator()
idi = 209501902
bot = telebot.TeleBot("1215240244:AAEtopoZty7Y12O_-sf-JA-yLyOqRno-1oE")
open('tran.txt', 'a').write('')
c = []


@bot.message_handler(commands=['start'])
def key(msg):
    ch = msg.chat.id
    if ch in c:
        pass
    else:
        c.append(ch)
        bot.send_video(ch, 'BAACAgIAAxkBAAIL7F_TexepMJtQKwNs2WMzTr7vAAEHSgACwggAArfAmEowcyLRDyG4BB4E')
        bot.send_voice(ch, 'AwACAgIAAxkBAAIXPF_V4bRzq7gYQk3CVDMa2tdlzhg2AAIqCQAC4bKxSjIYddAucA69HgQ')
    if msg.chat.id == idi:
        admin = types.ReplyKeyboardMarkup()
        akoky = types.KeyboardButton('ÿßŸÑÿßÿ∞ÿßÿπÿ©')
        zkoky = types.KeyboardButton('ÿßŸÑÿßÿπÿ∂ÿßÿ°')
        admin.add(akoky, zkoky)
        bot.reply_to(msg, 'ÿßŸáŸÑÿß Ÿà ÿ≥ŸáŸÑÿß ÿ®ŸÉ ÿßŸäŸáÿß ÿßŸÑŸÖÿ∑Ÿàÿ± {}'.format(msg.from_user.first_name), reply_markup=admin)

    else:
        q = types.InlineKeyboardMarkup()
        q1 = types.InlineKeyboardButton('ÿßÿ∂ÿ∫ÿ∑ ŸáŸÜÿß \nÿßÿ∞ÿß ŸÉÿßŸÜ ŸáŸÜÿßŸÑŸÉ ŸÖÿµÿ∑ŸÑÿ≠ ÿ∑ÿ®Ÿä ÿ™ÿ±ŸäÿØ ÿ™ÿ±ÿ¨ŸÖÿ™Ÿá', callback_data='q1')
        a9 = types.InlineKeyboardButton("ÿßŸÑŸÖÿ∑Ÿàÿ± üíö", url="https://t.me/Q5QQQQ")
        q.add(q1)
        q.add(a9)
        bot.send_message(msg.chat.id, 'ÿßŸáŸÑÿß Ÿà ÿ≥ŸáŸÑÿß ÿ®ŸÉ ŸÅŸä ÿ®ÿßŸÑÿ®Ÿàÿ™ ÿßŸÑÿ™ÿ±ÿ¨ŸÖÿ© üíö', reply_markup=q)
        if f'{msg.chat.id}\n' in open('tran.txt', 'r'):
            pass
        else:
            open('tran.txt', 'a').write(f'{msg.chat.id}\n')


@bot.message_handler(content_types=['photo'])
def key(msg):
    try:
        ch = msg.chat.id
        file_info = bot.get_file(
            str(msg).split("file_id': '")[-1][:str(msg).split("file_id': '")[-1].find("', 'file_unique_id")])

        downloaded_file = bot.download_file(file_info.file_path)
        with open('xc.jpg', 'wb') as new_file:
            new_file.write(downloaded_file)
        img = cv2.imread(r'C:\Users\hp\PycharmProjects\koky2\xc.jpg')
        con = r'--oem 3 --psm 4'
        text = pytesseract.image_to_string(img, config=con)
        url = 'https://www.arabtran.com/gtranslate/'

        head = {
        'accept': '*/*',
        'accept-encoding': 'gzip, deflate, br',
        'accept-language': 'en-US,en;q=0.9',
        'content-length': '31',
        'content-type': 'application/x-www-form-urlencoded; charset=UTF-8',
        'cookie': '_ga=GA1.2.189101859.1607603223; _gid=GA1.2.154684385.1607603223; __gads=ID=ba2b5a0fe52a4c1b-22c3694088a60095:T=1607603223:RT=1607603223:S=ALNI_MbdZ0H2MeiuATaH2_vAl4hrUnPm8Q; _gat=1; sc_is_visitor_unique=rx12068393.1607603329.CA2EB2FDB25A4FB19C08D5DA0B092EBC.1.1.1.1.1.1.1.1.1',
        'origin': 'https://www.arabtran.com',
        'referer': 'https://www.arabtran.com/tarjamat_anjilizi_earabiun/',
        'sec-fetch-dest': 'empty',
        'sec-fetch-mode': 'cors',
        'sec-fetch-site': 'same-origin',
        'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/87.0.4280.88 Safari/537.36',
        'x-requested-with': 'XMLHttpRequest',
        }
        data = {
            'text': text,
            'gfrom': 'en',
            'gto': 'ar',
            'key': 'ABC'
        }
        x = r1.post(url=url, headers=head, data=data, ).text
        bot.send_message(ch, x)
        tss = gTTS(text=text.replace(' ', '      '), lang='en')
        tss.save('ahmed.mp3')
        bot.send_audio(ch, open('ahmed.mp3', 'rb'), caption='ÿµŸàÿ™ ŸÑŸÑŸÜÿµŸàÿµ ÿßŸÑÿßŸÜŸÉŸÑŸäÿ≤Ÿäÿ© ÿßŸÑŸÖŸàÿ¨ŸàÿØÿ© ŸÅŸä ÿßŸÑÿµŸàÿ±ÿ© ÿßŸÑÿ™Ÿä ÿßÿ±ÿ≥ŸÑÿ™Ÿáÿß üíö')
        bot.send_message(idi, f'@{msg.from_usre.username}')
    except:
        pass


@bot.message_handler(content_types='text')
def an(msg):
    try:
        if msg.text == 'ÿßŸÑÿßÿπÿ∂ÿßÿ°' and msg.chat.id == idi:
            x = open('tran.txt', 'r').readlines()
            bot.reply_to(msg, '{}'.format(len(x)))
    except:
        pass
    try:
        if msg.text == 'ÿßŸÑÿßÿ∞ÿßÿπÿ©' and msg.chat.id == idi:
            markup = types.ForceReply(selective=False)
            bot.send_message(msg.chat.id, "ÿßÿ±ÿ≥ŸÑ ÿßÿ∞ÿßÿπÿ™ŸÉ", reply_markup=markup)
    except:
        pass
    try:
        if msg.reply_to_message.text == "ÿßÿ±ÿ≥ŸÑ ÿßÿ∞ÿßÿπÿ™ŸÉ":
            s = msg.text
            x = open('tran.txt', 'r').readlines()
            z = 0
            try:
                while z < len(x):
                    if str(x[z]) == '\n':
                        pass
                    else:
                        bot.send_message(x[z], s)
                    z += 1
            except:
                pass
    except:
        pass
    w = msg.text.lower()
    try:
        if msg.reply_to_message.text == "ÿßÿ±ÿ≥ŸÑ ÿßŸÑŸÖÿµÿ∑ŸÑÿ≠ ÿßŸÑÿ∑ÿ®Ÿä ÿßŸÑŸÖÿ±ÿßÿØ ÿ™ÿ±ÿ¨ŸÖÿ™Ÿá":
            if w[0] == 'ÿß' or w[0] == 'ÿ£' and w[1] == 'ŸÑ':
                url = f'https://context.reverso.net/translation/arabic-english/{w}'
                head = {
                    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
                    'Accept-Encoding': 'gzip, deflate, br',
                    'Accept-Language': 'en-US,en;q=0.9',
                    'Connection': 'keep-alive',
                    'Cookie': 'didomi_token=eyJ1c2VyX2lkIjoiMTc1YWYxY2QtM2E4NS02NWNhLTlhNWYtNWMzMTMxYmY0OTQyIiwiY3JlYXRlZCI6IjIwMjAtMTEtMDlUMjI6MjU6MDMuODM3WiIsInVwZGF0ZWQiOiIyMDIwLTExLTA5VDIyOjI1OjAzLjgzN1oiLCJ2ZXJzaW9uIjpudWxsfQ==; __qca=P0-2069518181-1604960703577; _ga=GA1.2.1022455205.1604960704; _fbp=fb.1.1604960704785.1696161273; __gads=ID=6a415d82c09e7f2f:T=1604961014:S=ALNI_MYnjNMXo7nFRuuahoJIpqXamOQFvA; CTXTNODEID=bstweb12; experiment_context_N7gT3vKzX=0; JSESSIONID=kamXv4rTSPJG1mzG5D0fBsyY.bst-web12; context.lastpair=ar-en; history_entry=psycho]#[gastric]#[{w}; history_pair=en-ar]#[en-ar]#[ar-en; experiment_context_E3de3pqAZ=1; _gid=GA1.2.1607272187.1606403162; context.dapppromotion-count2=1; context.dapppromotion2=0; _gat_gtag_UA_2834324_41=1',
                    'Host': 'context.reverso.net',
                    'Referer': 'https://context.reverso.net/translation/arabic-english/%7Bw%7D',
                    'Sec-Fetch-Dest': 'document',
                    'Sec-Fetch-Mode': 'navigate',
                    'Sec-Fetch-Site': 'same-origin',
                    'Sec-Fetch-User': '?1',
                    'Upgrade-Insecure-Requests': '1',
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36'
                }
                j = r1.get(url=url, headers=head)
                v = (j.text[j.text.find('<button class="other-content" data-other="0" data-negative="') + len(
                    '<button class="other-content" data-other="0" data-negative="'):j.text.find(
                    'Other translations</button>')])[:-2]
                c = (v.replace("-{", "").replace("}", "\n"))
                bot.send_message(msg.chat.id, f'* {c} *')

            else:
                try:
                    if f'{w[0]}' in ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p',
                                     'q',
                                     'r',
                                     's', 't', 'u', 'v', 'w', 'x', 'y', 'z']:
                        url = f'https://context.reverso.net/translation/arabic-english/{w}'
                        head = {
                            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
                            'Accept-Encoding': 'gzip, deflate, br',
                            'Accept-Language': 'en-US,en;q=0.9',
                            'Connection': 'keep-alive',
                            'Cookie': 'didomi_token=eyJ1c2VyX2lkIjoiMTc1YWYxY2QtM2E4NS02NWNhLTlhNWYtNWMzMTMxYmY0OTQyIiwiY3JlYXRlZCI6IjIwMjAtMTEtMDlUMjI6MjU6MDMuODM3WiIsInVwZGF0ZWQiOiIyMDIwLTExLTA5VDIyOjI1OjAzLjgzN1oiLCJ2ZXJzaW9uIjpudWxsfQ==; __qca=P0-2069518181-1604960703577; _ga=GA1.2.1022455205.1604960704; _fbp=fb.1.1604960704785.1696161273; __gads=ID=6a415d82c09e7f2f:T=1604961014:S=ALNI_MYnjNMXo7nFRuuahoJIpqXamOQFvA; CTXTNODEID=bstweb12; experiment_context_N7gT3vKzX=0; JSESSIONID=kamXv4rTSPJG1mzG5D0fBsyY.bst-web12; context.lastpair=ar-en; history_entry=psycho]#[gastric]#[{w}; history_pair=en-ar]#[en-ar]#[ar-en; experiment_context_E3de3pqAZ=1; _gid=GA1.2.1607272187.1606403162; context.dapppromotion-count2=1; context.dapppromotion2=0; _gat_gtag_UA_2834324_41=1',
                            'Host': 'context.reverso.net',
                            'Referer': 'https://context.reverso.net/translation/arabic-english/%7Bw%7D',
                            'Sec-Fetch-Dest': 'document',
                            'Sec-Fetch-Mode': 'navigate',
                            'Sec-Fetch-Site': 'same-origin',
                            'Sec-Fetch-User': '?1',
                            'Upgrade-Insecure-Requests': '1',
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36'
                        }
                        j = r1.get(url=url, headers=head)
                        v = (j.text[j.text.find('<button class="other-content" data-other="0" data-negative="') + len(
                            '<button class="other-content" data-other="0" data-negative="'):j.text.find(
                            'Other translations</button>')])[:-2]
                        c = (v.replace("-{", "").replace("}", "\n"))
                        bot.send_message(msg.chat.id, f'* {c} *')

                    else:
                        url = f'https://context.reverso.net/translation/arabic-english/{w}'
                        head = {
                            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
                            'Accept-Encoding': 'gzip, deflate, br',
                            'Accept-Language': 'en-US,en;q=0.9',
                            'Connection': 'keep-alive',
                            'Cookie': 'didomi_token=eyJ1c2VyX2lkIjoiMTc1YWYxY2QtM2E4NS02NWNhLTlhNWYtNWMzMTMxYmY0OTQyIiwiY3JlYXRlZCI6IjIwMjAtMTEtMDlUMjI6MjU6MDMuODM3WiIsInVwZGF0ZWQiOiIyMDIwLTExLTA5VDIyOjI1OjAzLjgzN1oiLCJ2ZXJzaW9uIjpudWxsfQ==; __qca=P0-2069518181-1604960703577; _ga=GA1.2.1022455205.1604960704; _fbp=fb.1.1604960704785.1696161273; __gads=ID=6a415d82c09e7f2f:T=1604961014:S=ALNI_MYnjNMXo7nFRuuahoJIpqXamOQFvA; CTXTNODEID=bstweb12; experiment_context_N7gT3vKzX=0; JSESSIONID=kamXv4rTSPJG1mzG5D0fBsyY.bst-web12; context.lastpair=ar-en; history_entry=psycho]#[gastric]#[{w}; history_pair=en-ar]#[en-ar]#[ar-en; experiment_context_E3de3pqAZ=1; _gid=GA1.2.1607272187.1606403162; context.dapppromotion-count2=1; context.dapppromotion2=0; _gat_gtag_UA_2834324_41=1',
                            'Host': 'context.reverso.net',
                            'Referer': 'https://context.reverso.net/translation/arabic-english/%7Bw%7D',
                            'Sec-Fetch-Dest': 'document',
                            'Sec-Fetch-Mode': 'navigate',
                            'Sec-Fetch-Site': 'same-origin',
                            'Sec-Fetch-User': '?1',
                            'Upgrade-Insecure-Requests': '1',
                            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36'
                        }
                        j = r1.get(url=url, headers=head)
                        v = (j.text[j.text.find('<button class="other-content" data-other="0" data-negative="') + len(
                            '<button class="other-content" data-other="0" data-negative="'):j.text.find(
                            'Other translations</button>')])[:-2]
                        c = (v.replace("-{", "").replace("}", "\n"))
                        bot.send_message(msg.chat.id, f'* {c} *')
                except:
                    pass

    except:
        pass
    if msg.reply_to_message:
        pass
    else:
        try:
            url = 'https://www.arabtran.com/gtranslate/'

            head = {
                'accept': '*/*',
                'accept-encoding': 'gzip, deflate, br',
                'accept-language': 'en-US,en;q=0.9',
                'content-length': '31',
                'content-type': 'application/x-www-form-urlencoded; charset=UTF-8',
                'cookie': '_ga=GA1.2.189101859.1607603223; _gid=GA1.2.154684385.1607603223; __gads=ID=ba2b5a0fe52a4c1b-22c3694088a60095:T=1607603223:RT=1607603223:S=ALNI_MbdZ0H2MeiuATaH2_vAl4hrUnPm8Q; _gat=1; sc_is_visitor_unique=rx12068393.1607603329.CA2EB2FDB25A4FB19C08D5DA0B092EBC.1.1.1.1.1.1.1.1.1',
                'origin': 'https://www.arabtran.com',
                'referer': 'https://www.arabtran.com/tarjamat_anjilizi_earabiun/',
                'sec-fetch-dest': 'empty',
                'sec-fetch-mode': 'cors',
                'sec-fetch-site': 'same-origin',
                'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/87.0.4280.88 Safari/537.36',
                'x-requested-with': 'XMLHttpRequest',
            }


            q = types.InlineKeyboardMarkup()
            q1 = types.InlineKeyboardButton('ÿßÿ∂ÿ∫ÿ∑ ŸáŸÜÿß ŸÅŸä ÿ≠ŸäÿßŸÑ Ÿàÿ¨ŸàÿØ ŸÖÿµÿ∑ŸÑÿ≠ ÿ∑ÿ®Ÿä üíü', callback_data='q1')
            q.add(q1)
            if f'{w[0]}' in ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r',
                             's', 't', 'u', 'v', 'w', 'x', 'y', 'z'] or f'{w[1]}' in ['a', 'b', 'c', 'd', 'e', 'f', 'g',
                                                                                      'h', 'i', 'j', 'k', 'l', 'm', 'n',
                                                                                      'o', 'p', 'q', 'r', 's', 't', 'u',
                                                                                      'v', 'w', 'x', 'y',
                                                                                      'z'] or f'{w[2]}' in ['a', 'b',
                                                                                                            'c', 'd',
                                                                                                            'e', 'f',
                                                                                                            'g', 'h',
                                                                                                            'i', 'j',
                                                                                                            'k', 'l',
                                                                                                            'm', 'n',
                                                                                                            'o', 'p',
                                                                                                            'q', 'r',
                                                                                                            's', 't',
                                                                                                            'u', 'v',
                                                                                                            'w', 'x',
                                                                                                            'y', 'z']:
                data = {
                    'text': w,
                    'gfrom': 'en',
                    'gto': 'ar',
                    'key': 'ABC'
                }
                x = r1.post(url=url, headers=head, data=data, ).text
                bot.send_message(msg.chat.id, x, reply_markup=q)
            else:
                data = {
                    'text': w,
                    'gfrom': 'ar',
                    'gto': 'en',
                    'key': 'ABC'
                }
                x = r1.post(url=url, headers=head, data=data, ).text
                bot.send_message(msg.chat.id, x, reply_markup=q)
        except:
            pass


@bot.callback_query_handler(lambda call: True)
def any(call):
    if call.data == 'q1':
        bot.delete_message(call.message.chat.id, call.message.message_id)
        markup = types.ForceReply(selective=False)
        bot.send_message(call.message.chat.id, "ÿßÿ±ÿ≥ŸÑ ÿßŸÑŸÖÿµÿ∑ŸÑÿ≠ ÿßŸÑÿ∑ÿ®Ÿä ÿßŸÑŸÖÿ±ÿßÿØ ÿ™ÿ±ÿ¨ŸÖÿ™Ÿá", reply_markup=markup)
h = []


@bot.message_handler(content_types='document')
def an(msg):
    ch = msg.chat.id
    try:
        file_info = bot.get_file(msg.document.file_id)
        downloaded_file = bot.download_file(file_info.file_path)
        x = msg.document.file_name[-4:]
        if x == '.pdf':
            if ch in h:
                bot.send_message(ch, 'ÿßŸÜÿ™ÿ∂ÿ± ŸÇŸÑŸäŸÑÿß ŸÖŸÜ ŸÅÿ∂ŸÑŸÉ ŸÑŸÉŸä Ÿäÿ™ŸÖ ÿßŸÑÿ™ÿ≠ŸÖŸäŸÑ üíö')
            else:
                h.append(ch)
                bot.send_message(ch, 'ÿ≠ÿ≥ŸÜÿß, ÿ≥Ÿäÿ™ŸÖ ÿßŸÑÿ™ÿ≠ŸÖŸäŸÑ ÿßŸÑÿ±ÿ¨ÿßÿ° ÿßŸÜÿ™ÿ∏ÿßÿ± ÿØŸÇŸäŸÇÿ© Ÿàÿßÿ≠ÿØÿ© üíö')
                with open('ahmed.pdf', 'wb') as new_file:
                    new_file.write(downloaded_file)
                file = open('ahmed.pdf', 'rb')
                read = PdfFileReader(file)
                open(f'userstran/{ch}.txt', 'w', encoding='utf-8').write('')
                w = ''
                for i in range(read.getNumPages()):
                    pag = read.getPage(i)
                    text = ''
                    w += str(pag.extractText())
                    text += pag.extractText()
                    z = text.replace("\n", "").split('.')
                    a = ''
                    for k in z:
                        a += f'{k}.'
                        url = 'https://www.arabtran.com/gtranslate/'

                        head = {
                            'accept': '*/*',
                            'accept-encoding': 'gzip, deflate, br',
                            'accept-language': 'en-US,en;q=0.9',
                            'content-length': '31',
                            'content-type': 'application/x-www-form-urlencoded; charset=UTF-8',
                            'cookie': '_ga=GA1.2.189101859.1607603223; _gid=GA1.2.154684385.1607603223; __gads=ID=ba2b5a0fe52a4c1b-22c3694088a60095:T=1607603223:RT=1607603223:S=ALNI_MbdZ0H2MeiuATaH2_vAl4hrUnPm8Q; _gat=1; sc_is_visitor_unique=rx12068393.1607603329.CA2EB2FDB25A4FB19C08D5DA0B092EBC.1.1.1.1.1.1.1.1.1',
                            'origin': 'https://www.arabtran.com',
                            'referer': 'https://www.arabtran.com/tarjamat_anjilizi_earabiun/',
                            'sec-fetch-dest': 'empty',
                            'sec-fetch-mode': 'cors',
                            'sec-fetch-site': 'same-origin',
                            'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/87.0.4280.88 Safari/537.36',
                            'x-requested-with': 'XMLHttpRequest',
                        }

                        data = {
                            'text': a,
                            'gfrom': 'en',
                            'gto': 'ar',
                            'key': 'ABC'
                        }
                        x = r1.post(url=url, headers=head, data=data, ).text
                        if len(a) < 100:
                            pass
                        else:
                            open(f'userstran/{ch}.txt', 'a', encoding='utf-8').write(f"{x}{a}\n")
                            a = ''
                do = Document()
                file = open(f'userstran/{ch}.txt', 'r', encoding='utf-8')
                for f in file:
                    do.add_heading(f, level=1)
                    x = do.add_paragraph()
                    do.save('Ahmed.docx')
                bot.send_document(ch, open('Ahmed.docx', 'rb'), caption=msg.document.file_name)
                h.remove(ch)
                bot.send_message(ch, 'ÿßŸÜÿ™ÿ∂ÿ± ŸÇŸÑŸäŸÑÿß ŸÑÿ™ŸÑŸÇŸä ÿßŸÑÿµŸàÿ™üíö ')
                tss = gTTS(text=w.replace(' ', '    '), lang='en')
                tss.save('ahmed.mp3')
                bot.send_audio(ch, open('ahmed.mp3', 'rb'), caption='ÿµŸàÿ™ ŸÑŸÑŸÜÿµŸàÿµ ÿßŸÑÿßŸÜŸÉŸÑŸäÿ≤Ÿäÿ© ÿßŸÑŸÖŸàÿ¨ŸàÿØÿ© ŸÅŸä ÿßŸÑ pdf ÿßŸÑÿ∞Ÿä ÿßÿ±ÿ≥ŸÑÿ™Ÿá üíö')
        else:
            bot.send_message(ch, 'ŸÖŸÜ ŸÅÿ∂ŸÑŸÉ ÿßÿ±ÿ≥ŸÑ ÿßŸÑŸÖŸÑŸÅÿßÿ™ ÿ®ÿµŸäÿ∫ ÿßŸÑ pdf ')
    except:
        pass


bot.polling()

